import requests
import os
import re
import time
import random
from fake_useragent import UserAgent
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
url_list = 'https://www.83ks.org/book/10873/'
headers={'User-Agent':UserAgent().random}

dict = {}
link_re = re.compile(r'<a href="(?P<link>[\S]*?)" title="总裁在上[\s\S]*?>(?P<title>[\s\S]*?)</a>')
content_re = re.compile(r'<div class="content" id="content">[\s\S]*?<p>(?P<content>[\s\S]*?)</div>')
resp_list = requests.get(url=url_list, headers=headers)
print(resp_list)
resp_list_text = resp_list.text
result_re = link_re.finditer(resp_list_text)
for i in result_re:
    link = i.group('link').strip()
    title = i.group('title').strip()
    dict[title] = link
    
num = 1
content = ''
'''
url_chapter = 'https://www.83ks.org' + dict[list(dict.keys())[0]]
resp_content = requests.get(url = url_chapter, headers=headers)
resp_content.encoding = resp_content.apparent_encoding
'''
for title in dict:
    url_chapter = 'https://www.83ks.org' + dict[title]
    resp_content = requests.get(url = url_chapter, headers=headers)
    print(resp_content)
    result = content_re.finditer(resp_content.text)
    content += title + '\n'
    for i in result:
        content += i.group('content').strip()
    content += '\n'
    content = content.replace('<p>', '')
    content = content.replace('</p>', '\n')
    sleeptime = random.randint(1, 10)
    print(num)
    time.sleep(sleeptime)
    num += 1
'''
with open('总裁在上.txt', 'w', encoding=resp_content.encoding) as f:
    f.write(content)
'''

doc = Document()
doc.styles['Normal'].font.name = u'宋体'
for line in content.splitlines():
    if line.strip():  # 只添加非空行
        paragraph = doc.add_paragraph(line.strip())
        run = paragraph.runs[0]
        run.font.name = 'SimSun'  # 设置字体为宋体
        # 下面两行确保中文字体能被识别
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)  # 设置字体大小为12pt
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 左对齐
    else:
        paragraph = doc.add_paragraph(line.strip())

# 保存文档
doc.save('C:/Document/总裁在上.docx')



